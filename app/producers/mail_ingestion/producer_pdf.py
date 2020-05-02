from imap_server import HyperplanImapServer 
from PyPDF2 import PdfFileReader
import time
import base64
import persistqueue
import os
import docx

class Producer:
    
    def __init__(self):
        
        self.list_classified_mails = []
        self.q = persistqueue.SQLiteQueue(os.path.join(os.path.dirname(__file__), "../../queue"), auto_commit=True)
        self.imapServerConnection()
        
    def imapServerConnection(self):
        
        hyperplan = HyperplanImapServer()
        imap = hyperplan.imapConnection()
        hyperplan.storeMails(imap, True)
        informations = hyperplan.fetchMails()
        self.attachmentToText(informations)
        print('Searching mails...\n')
        
        try:
            while True:
                if (hyperplan.storeMails(imap)):
                    informations = hyperplan.fetchMails()
                    self.attachmentToText(informations)
                    print('Searching mails...\n')
                    
                time.sleep(5)
        except KeyboardInterrupt: # Ctrl + c
            pass            
            
        hyperplan.imapDisconnection(imap)
        
    def feedFIFO(self, id, text):
        
        self.q.put({"id": id, "clean_text": text})
    
    def pdfToText(self, attachmentName, fileContent):
        
        open('./utils/files/' + attachmentName, 'wb').write(base64.b64decode(fileContent))
                                
        with open('./utils/files/' + attachmentName, 'rb') as f:
            pdfReader = PdfFileReader(f)
            pageObject = ''
            
            for page in pdfReader.pages:
                pageObject += page.extractText()
            
            attachment_text = ''
            
            for line in pageObject.replace('\n', '').split('  '):
                attachment_text += line + '\n'
            
        os.remove('./utils/files/' + attachmentName)
        
        return attachment_text

    def docxToText(self, attachmentName, fileContent):
        
        open('./utils/files/' + attachmentName, 'wb').write(base64.b64decode(fileContent))

        doc = docx.Document('./utils/files/' + attachmentName)
        attachment_text = ''

        for para in doc.paragraphs:
            attachment_text += ' ' + para.text

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        attachment_text += ' ' + para.text
            
        os.remove('./utils/files/' + attachmentName)
        
        return attachment_text
                                
    def attachmentToText(self, informations):
        
        for mail in informations:
            if (mail[0] not in self.list_classified_mails):
                for info in mail:
                    if (mail.index(info) == 0):
                        mail_id = info
                        print('Mail ID : %s\n' % mail_id)
                        self.list_classified_mails.append(mail_id)
                        
                    if (mail.index(info) == 5):
                        attachments = info.split(',')

                        for i in range(0, len(attachments), 2):
                            attachmentName = attachments[i]
                            
                            if attachmentName:
                                fileContent = attachments[i + 1]
                                attachment_text = ''
                                
                                print('Nom de la pièce jointe : %s\n' % attachmentName)
                                
                                if attachmentName.endswith('.pdf'):                                
                                    attachment_text = self.pdfToText(attachmentName, fileContent)
                                    
                                elif attachmentName.endswith('.docx'):
                                    attachment_text = self.docxToText(attachmentName, fileContent)
                                    
                                self.feedFIFO(mail_id, attachment_text)

if __name__ == "__main__":
    producer = Producer()
    